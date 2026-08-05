import json
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
PIPELINE = REPO_ROOT / "plugins" / "office-toolkit-windows" / "skills" / "pdf-scan-to-docx" / "scripts" / "pipeline.py"


def run_pipeline(*args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(PIPELINE), *map(str, args)],
        capture_output=True,
        text=True,
        check=False,
    )


def make_pdf(path: Path, pages: int = 2) -> None:
    pdf = fitz.open()
    for number in range(1, pages + 1):
        page = pdf.new_page()
        page.insert_text((72, 72), f"Trang {number}")
    pdf.save(path)
    pdf.close()


class PdfPipelineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp_dir.cleanup)
        self.workspace = Path(self.temp_dir.name) / "job"
        prepared = run_pipeline("prepare", self.workspace)
        self.assertEqual(prepared.returncode, 0, prepared.stderr)

    def test_render_preprocess_resume_merge_and_export_keep_page_checkpoints(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source)

        rendered = run_pipeline("render", self.workspace, source)
        self.assertEqual(rendered.returncode, 0, rendered.stderr)
        self.assertTrue((self.workspace / "02.process" / "page-0001.png").is_file())
        self.assertTrue((self.workspace / "02.process" / "page-0002.png").is_file())

        preprocessed = run_pipeline("preprocess", self.workspace)
        self.assertEqual(preprocessed.returncode, 0, preprocessed.stderr)
        self.assertTrue((self.workspace / "02.process" / "preprocessed" / "page-0001.png").is_file())

        resumed = run_pipeline("render", self.workspace, source)
        self.assertEqual(resumed.returncode, 0, resumed.stderr)
        self.assertIn("rendered=0", resumed.stdout)

        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0002.md").write_text("Trang hai", encoding="utf-8")
        (ocr / "page-0001.md").write_text("Trang mot", encoding="utf-8")

        merged = run_pipeline("merge", self.workspace)
        self.assertEqual(merged.returncode, 0, merged.stderr)
        markdown = self.workspace / "03.output" / "mau.md"
        self.assertEqual(markdown.read_text(encoding="utf-8"), "Trang mot\n\nTrang hai\n")

        exported = run_pipeline("export", self.workspace)
        self.assertEqual(exported.returncode, 0, exported.stderr)
        document = Document(self.workspace / "03.output" / "mau.docx")
        self.assertEqual([paragraph.text for paragraph in document.paragraphs], ["Trang mot", "Trang hai"])

    def test_merge_rejects_a_missing_page_checkpoint(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("only one", encoding="utf-8")

        result = run_pipeline("merge", self.workspace)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("missing OCR checkpoint", result.stderr)
        self.assertFalse((self.workspace / "03.output" / "mau.md").exists())

    def test_render_rejects_a_pdf_outside_input(self) -> None:
        outside = Path(self.temp_dir.name) / "outside.pdf"
        make_pdf(outside, pages=1)

        result = run_pipeline("render", self.workspace, outside)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("01.input", result.stderr)
        self.assertFalse((self.workspace / "02.process" / "page-0001.png").exists())

    def test_commands_reject_a_workspace_without_the_marker(self) -> None:
        wrong = Path(self.temp_dir.name) / "wrong"
        for name in ("01.input", "02.process", "03.output"):
            (wrong / name).mkdir(parents=True, exist_ok=True)

        result = run_pipeline("preprocess", wrong)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("workspace marker", result.stderr)

    def test_cleanup_needs_confirmation_and_completed_outputs(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("done", encoding="utf-8")
        self.assertEqual(run_pipeline("merge", self.workspace).returncode, 0)
        self.assertEqual(run_pipeline("export", self.workspace).returncode, 0)
        output = self.workspace / "03.output"
        keep = self.workspace / "keep.txt"
        keep.write_text("keep", encoding="utf-8")

        refused = run_pipeline("cleanup", self.workspace)
        self.assertNotEqual(refused.returncode, 0)
        self.assertTrue(source.exists())
        self.assertTrue((self.workspace / "02.process").exists())

        cleaned = run_pipeline("cleanup", self.workspace, "--confirm")
        self.assertEqual(cleaned.returncode, 0, cleaned.stderr)
        self.assertFalse((self.workspace / "01.input").exists())
        self.assertFalse((self.workspace / "02.process").exists())
        self.assertTrue((output / "mau.md").is_file())
        self.assertTrue((output / "mau.docx").is_file())
        self.assertTrue(keep.is_file())

    def test_merge_rejects_a_manifest_stem_that_escapes_output(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("safe", encoding="utf-8")
        manifest = self.workspace / "02.process" / "source.json"
        info = json.loads(manifest.read_text(encoding="utf-8"))
        info["stem"] = "../escaped"
        manifest.write_text(json.dumps(info), encoding="utf-8")

        result = run_pipeline("merge", self.workspace)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("output stem", result.stderr)
        self.assertFalse((self.workspace / "escaped.md").exists())

    def test_resume_rejects_a_source_pdf_changed_after_render(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=2)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        replacement = Path(self.temp_dir.name) / "replacement.pdf"
        make_pdf(replacement, pages=1)
        replacement.replace(source)

        result = run_pipeline("render", self.workspace, source)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("source changed", result.stderr)
        self.assertTrue((self.workspace / "02.process" / "page-0002.png").is_file())

    def test_merge_rejects_a_missing_rendered_page_from_manifest_count(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=2)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        (self.workspace / "02.process" / "page-0002.png").unlink()
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("one", encoding="utf-8")

        result = run_pipeline("merge", self.workspace)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("rendered page set", result.stderr)

    def test_merge_rejects_an_extra_rendered_page_outside_manifest_count(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=2)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        process = self.workspace / "02.process"
        shutil.copyfile(process / "page-0001.png", process / "page-0003.png")
        ocr = process / "ocr"
        ocr.mkdir()
        for number in (1, 2, 3):
            (ocr / f"page-{number:04d}.md").write_text(str(number), encoding="utf-8")

        result = run_pipeline("merge", self.workspace)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("rendered page set", result.stderr)

    def test_merge_rejects_a_reparse_rendered_page_entry(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        page = self.workspace / "02.process" / "page-0001.png"
        page.unlink()
        target = Path(self.temp_dir.name) / "target"
        target.mkdir()
        junction = subprocess.run(
            ["cmd.exe", "/d", "/c", "mklink", "/J", str(page), str(target)],
            capture_output=True,
            text=True,
            check=False,
        )
        if junction.returncode != 0:
            self.skipTest(f"junction unavailable: {junction.stderr}")
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("one", encoding="utf-8")

        result = run_pipeline("merge", self.workspace)

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("reparse", result.stderr)

    def test_cleanup_rejects_a_stale_output_pair_not_bound_to_current_job(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        output = self.workspace / "03.output"
        (output / "stale.md").write_text("stale", encoding="utf-8")
        (output / "stale.docx").write_bytes(b"stale")

        result = run_pipeline("cleanup", self.workspace, "--confirm")

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("current job", result.stderr)
        self.assertTrue(source.is_file())
        self.assertTrue((self.workspace / "02.process").is_dir())

    def test_cleanup_rejects_a_changed_source_after_outputs_exist(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("done", encoding="utf-8")
        self.assertEqual(run_pipeline("merge", self.workspace).returncode, 0)
        self.assertEqual(run_pipeline("export", self.workspace).returncode, 0)
        replacement = Path(self.temp_dir.name) / "replacement.pdf"
        make_pdf(replacement, pages=2)
        replacement.replace(source)

        result = run_pipeline("cleanup", self.workspace, "--confirm")

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("source changed", result.stderr)
        self.assertTrue(source.is_file())

    def test_cleanup_rejects_an_output_changed_after_export(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("done", encoding="utf-8")
        self.assertEqual(run_pipeline("merge", self.workspace).returncode, 0)
        self.assertEqual(run_pipeline("export", self.workspace).returncode, 0)
        markdown = self.workspace / "03.output" / "mau.md"
        markdown.write_text("changed", encoding="utf-8")

        result = run_pipeline("cleanup", self.workspace, "--confirm")

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("output changed", result.stderr)
        self.assertTrue(source.is_file())

    def test_cleanup_rejects_a_docx_changed_after_export(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text("done", encoding="utf-8")
        self.assertEqual(run_pipeline("merge", self.workspace).returncode, 0)
        self.assertEqual(run_pipeline("export", self.workspace).returncode, 0)
        document = self.workspace / "03.output" / "mau.docx"
        document.write_bytes(b"changed")

        result = run_pipeline("cleanup", self.workspace, "--confirm")

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("output changed", result.stderr)
        self.assertTrue(source.is_file())

    def test_export_preserves_markdown_headings_lists_and_table(self) -> None:
        source = self.workspace / "01.input" / "mau.pdf"
        make_pdf(source, pages=1)
        self.assertEqual(run_pipeline("render", self.workspace, source).returncode, 0)
        ocr = self.workspace / "02.process" / "ocr"
        ocr.mkdir()
        (ocr / "page-0001.md").write_text(
            "# Title\n\n## Section\n\n- first\n- second\n\n1. one\n2. two\n\n| Left | Right |\n| --- | --- |\n| A | B |",
            encoding="utf-8",
        )
        self.assertEqual(run_pipeline("merge", self.workspace).returncode, 0)

        exported = run_pipeline("export", self.workspace)

        self.assertEqual(exported.returncode, 0, exported.stderr)
        document = Document(self.workspace / "03.output" / "mau.docx")
        paragraph_styles = {paragraph.text: paragraph.style.name for paragraph in document.paragraphs}
        self.assertEqual(paragraph_styles["Title"], "Heading 1")
        self.assertEqual(paragraph_styles["Section"], "Heading 2")
        self.assertEqual(paragraph_styles["first"], "List Bullet")
        self.assertEqual(paragraph_styles["one"], "List Number")
        self.assertEqual([[cell.text for cell in row.cells] for row in document.tables[0].rows], [["Left", "Right"], ["A", "B"]])

    def test_cleanup_rejects_a_reparse_process_target(self) -> None:
        outside = Path(self.temp_dir.name) / "outside"
        outside.mkdir()
        process = self.workspace / "02.process"
        process.rmdir()
        try:
            process.symlink_to(outside, target_is_directory=True)
        except OSError as error:
            junction = subprocess.run(
                ["cmd.exe", "/d", "/c", "mklink", "/J", str(process), str(outside)],
                capture_output=True,
                text=True,
                check=False,
            )
            if junction.returncode != 0:
                self.skipTest(f"symlink unavailable: {error}; junction unavailable: {junction.stderr}")
        output = self.workspace / "03.output"
        (output / "mau.md").write_text("done", encoding="utf-8")
        (output / "mau.docx").write_bytes(b"done")

        result = run_pipeline("cleanup", self.workspace, "--confirm")

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("reparse", result.stderr)
        self.assertTrue(outside.is_dir())


if __name__ == "__main__":
    unittest.main()
