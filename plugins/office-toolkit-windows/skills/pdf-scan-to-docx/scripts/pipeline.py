"""Local, checkpointed PDF scan preparation and DOCX export."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import sys
from pathlib import Path

import fitz
from docx import Document
from PIL import Image, ImageFilter, ImageOps


MARKER = ".pdf-scan-to-docx-workspace"
WORKSPACE_DIRS = ("01.input", "02.process", "03.output")
REPARSE_POINT = 0x400


class PipelineError(RuntimeError):
    pass


def is_reparse(path: Path) -> bool:
    if path.is_symlink():
        return True
    attributes = getattr(os.lstat(path), "st_file_attributes", 0)
    return bool(attributes & REPARSE_POINT)


def require_real(path: Path, kind: str) -> Path:
    if not path.exists() or is_reparse(path):
        raise PipelineError(f"{kind} is missing or a symlink/reparse target: {path}")
    return path.resolve(strict=True)


def workspace_paths(workspace: str | Path) -> dict[str, Path]:
    root = require_real(Path(workspace), "workspace")
    if not root.is_dir():
        raise PipelineError(f"workspace is not a directory: {root}")
    marker = root / MARKER
    if not marker.is_file() or is_reparse(marker):
        raise PipelineError(f"workspace marker is missing or unsafe: {marker}")
    paths = {"root": root}
    for name in WORKSPACE_DIRS:
        child = root / name
        if not child.is_dir() or is_reparse(child):
            raise PipelineError(f"workspace child is missing or a symlink/reparse target: {child}")
        paths[name] = child
    return paths


def require_direct_file(path: str | Path, directory: Path, label: str) -> Path:
    candidate = Path(path)
    resolved = require_real(candidate, label)
    if not resolved.is_file() or resolved.parent != directory:
        raise PipelineError(f"{label} must be a direct file inside {directory}")
    return resolved


def safe_child(directory: Path, name: str, required: bool = False) -> Path:
    if not isinstance(name, str) or not name or Path(name).name != name or name in {".", ".."}:
        raise PipelineError(f"unsafe direct child name: {name!r}")
    child = directory / name
    try:
        child.resolve(strict=False).relative_to(directory.resolve(strict=True))
    except ValueError as error:
        raise PipelineError(f"path escapes its allowed directory: {child}") from error
    if child.exists() and is_reparse(child):
        raise PipelineError(f"symlink/reparse target is not allowed: {child}")
    if required and not child.is_file():
        raise PipelineError(f"required output is missing: {child}")
    return child


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_source_info(paths: dict[str, Path], info: dict[str, object]) -> None:
    manifest = safe_child(paths["02.process"], "source.json")
    temporary = safe_child(paths["02.process"], "source.json.tmp")
    temporary.write_text(json.dumps(info, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(manifest)


def source_info(paths: dict[str, Path]) -> dict[str, object]:
    manifest = safe_child(paths["02.process"], "source.json", required=True)
    try:
        info = json.loads(manifest.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as error:
        raise PipelineError(f"source manifest is invalid: {manifest}") from error
    stem = info.get("stem")
    if not isinstance(stem, str) or not stem or Path(stem).name != stem or stem in {".", ".."}:
        raise PipelineError(f"source manifest has an unsafe output stem: {manifest}")
    source_name = info.get("source_name")
    if not isinstance(source_name, str) or Path(source_name).name != source_name or not source_name.lower().endswith(".pdf"):
        raise PipelineError(f"source manifest has an invalid source name: {manifest}")
    source_hash = info.get("source_sha256")
    if not isinstance(source_hash, str) or len(source_hash) != 64:
        raise PipelineError(f"source manifest has an invalid source hash: {manifest}")
    page_count = info.get("page_count")
    if not isinstance(page_count, int) or isinstance(page_count, bool) or page_count < 1:
        raise PipelineError(f"source manifest has an invalid page count: {manifest}")
    source = safe_child(paths["01.input"], source_name, required=True)
    if file_sha256(source) != source_hash:
        raise PipelineError(f"source changed after render: {source}")
    return info


def prepare(workspace: str | Path) -> None:
    root = Path(workspace)
    if root.exists():
        if is_reparse(root):
            raise PipelineError(f"workspace is a symlink/reparse target: {root}")
        if any(root.iterdir()):
            workspace_paths(root)
            return
    root.mkdir(parents=True, exist_ok=True)
    (root / MARKER).write_text("pdf-scan-to-docx\n", encoding="utf-8")
    for name in WORKSPACE_DIRS:
        (root / name).mkdir(exist_ok=True)


def render(workspace: str | Path, pdf: str | Path) -> int:
    paths = workspace_paths(workspace)
    source = require_direct_file(pdf, paths["01.input"], "PDF source")
    if source.suffix.lower() != ".pdf":
        raise PipelineError(f"PDF source must end in .pdf: {source}")
    source_hash = file_sha256(source)
    document = fitz.open(source)
    try:
        page_count = len(document)
        if page_count < 1:
            raise PipelineError(f"PDF source has no pages: {source}")
        manifest = safe_child(paths["02.process"], "source.json")
        if manifest.exists():
            existing = source_info(paths)
            if (
                existing["source_name"] != source.name
                or existing["source_sha256"] != source_hash
                or existing["page_count"] != page_count
            ):
                raise PipelineError(f"source changed after render: {source}")
        elif any(paths["02.process"].glob("page-????.png")) or safe_child(paths["02.process"], "ocr").exists():
            raise PipelineError("rendered/OCR artifacts exist without a source manifest")
        else:
            write_source_info(
                paths,
                {
                    "source_name": source.name,
                    "source_sha256": source_hash,
                    "stem": source.stem,
                    "page_count": page_count,
                },
            )
        rendered = 0
        for number, page in enumerate(document, start=1):
            target = safe_child(paths["02.process"], f"page-{number:04d}.png")
            if target.exists():
                continue
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            pixmap.save(target)
            rendered += 1
    finally:
        document.close()
    return rendered


def preprocess(workspace: str | Path) -> int:
    paths = workspace_paths(workspace)
    page_numbers = expected_pages(paths)
    destination = safe_child(paths["02.process"], "preprocessed")
    destination.mkdir(exist_ok=True)
    processed = 0
    for number in page_numbers:
        source = require_direct_file(
            safe_child(paths["02.process"], f"page-{number:04d}.png", required=True),
            paths["02.process"],
            "rendered page",
        )
        target = safe_child(destination, source.name)
        if target.exists():
            continue
        with Image.open(source) as image:
            prepared = ImageOps.autocontrast(image.convert("RGB")).filter(ImageFilter.SHARPEN)
            prepared.save(target, "PNG")
        processed += 1
    return processed


def expected_pages(paths: dict[str, Path]) -> list[int]:
    info = source_info(paths)
    process = paths["02.process"]
    actual: set[int] = set()
    for page in process.glob("page-????.png"):
        if is_reparse(page):
            raise PipelineError(f"symlink/reparse target is not allowed: {page}")
        if not page.is_file():
            raise PipelineError(f"rendered page is not a file: {page}")
        actual.add(int(page.stem.rsplit("-", 1)[1]))
    expected = set(range(1, int(info["page_count"]) + 1))
    if actual != expected:
        raise PipelineError(
            f"rendered page set does not match manifest page_count: expected {sorted(expected)}, got {sorted(actual)}"
        )
    return sorted(expected)


def merge(workspace: str | Path) -> Path:
    paths = workspace_paths(workspace)
    info = source_info(paths)
    ocr = safe_child(paths["02.process"], "ocr")
    if not ocr.is_dir() or is_reparse(ocr):
        raise PipelineError(f"OCR checkpoint directory is missing or unsafe: {ocr}")
    checkpoints: list[Path] = []
    for number in expected_pages(paths):
        checkpoint = safe_child(ocr, f"page-{number:04d}.md")
        if not checkpoint.is_file():
            raise PipelineError(f"missing OCR checkpoint: {checkpoint}")
        if is_reparse(checkpoint):
            raise PipelineError(f"symlink/reparse target is not allowed: {checkpoint}")
        checkpoints.append(checkpoint)
    output = safe_child(paths["03.output"], f"{info['stem']}.md")
    output.write_text("\n\n".join(item.read_text(encoding="utf-8").strip() for item in checkpoints) + "\n", encoding="utf-8")
    info.update(markdown_name=output.name, markdown_sha256=file_sha256(output))
    info.pop("docx_name", None)
    info.pop("docx_sha256", None)
    write_source_info(paths, info)
    return output


def markdown_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_to_document(document: Document, markdown: str) -> None:
    # ponytail: block-level Markdown only; use Pandoc if inline/complex fidelity becomes a requirement.
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            document.add_heading(heading.group(2).strip(), level=len(heading.group(1)))
            index += 1
            continue
        if re.match(r"^[-*+]\s+", line):
            document.add_paragraph(re.sub(r"^[-*+]\s+", "", line), style="List Bullet")
            index += 1
            continue
        if re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines):
            separator = markdown_cells(lines[index + 1])
            if separator and all(re.fullmatch(r":?-{3,}:?", cell) for cell in separator):
                rows = [markdown_cells(line)]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append(markdown_cells(lines[index]))
                    index += 1
                width = len(rows[0])
                if any(len(row) != width for row in rows):
                    raise PipelineError("Markdown table rows have inconsistent column counts")
                table = document.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for column_index, value in enumerate(row):
                        table.cell(row_index, column_index).text = value
                continue
        document.add_paragraph(line)
        index += 1


def export_docx(workspace: str | Path) -> Path:
    paths = workspace_paths(workspace)
    info = source_info(paths)
    markdown = safe_child(paths["03.output"], f"{info['stem']}.md", required=True)
    if info.get("markdown_name") != markdown.name:
        raise PipelineError("Markdown output is not bound to the current job; run merge first")
    if info.get("markdown_sha256") != file_sha256(markdown):
        raise PipelineError(f"output changed after merge: {markdown}")
    document = Document()
    add_markdown_to_document(document, markdown.read_text(encoding="utf-8"))
    output = safe_child(paths["03.output"], f"{info['stem']}.docx")
    document.save(output)
    info.update(docx_name=output.name, docx_sha256=file_sha256(output))
    write_source_info(paths, info)
    return output


def reject_reparse_tree(directory: Path) -> None:
    with os.scandir(directory) as entries:
        for entry in entries:
            child = Path(entry.path)
            if is_reparse(child):
                raise PipelineError(f"symlink/reparse target is not allowed: {child}")
            if entry.is_dir(follow_symlinks=False):
                reject_reparse_tree(child)


def cleanup(workspace: str | Path, confirmed: bool) -> None:
    if not confirmed:
        raise PipelineError("cleanup requires --confirm")
    paths = workspace_paths(workspace)
    info = source_info(paths)
    expected_names = {
        "markdown_name": f"{info['stem']}.md",
        "docx_name": f"{info['stem']}.docx",
    }
    if any(info.get(key) != name for key, name in expected_names.items()):
        raise PipelineError("required outputs are not bound to the current job")
    for name_key, hash_key in (("markdown_name", "markdown_sha256"), ("docx_name", "docx_sha256")):
        output = safe_child(paths["03.output"], str(info[name_key]), required=True)
        if info.get(hash_key) != file_sha256(output):
            raise PipelineError(f"output changed after export: {output}")
    for name in ("01.input", "02.process"):
        reject_reparse_tree(paths[name])
    for name in ("01.input", "02.process"):
        shutil.rmtree(paths[name])


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)
    for command in ("prepare", "preprocess", "merge", "export"):
        subparser = subparsers.add_parser(command)
        subparser.add_argument("workspace")
    render_parser = subparsers.add_parser("render")
    render_parser.add_argument("workspace")
    render_parser.add_argument("pdf")
    cleanup_parser = subparsers.add_parser("cleanup")
    cleanup_parser.add_argument("workspace")
    cleanup_parser.add_argument("--confirm", action="store_true")
    args = parser.parse_args()
    try:
        if args.command == "prepare":
            prepare(args.workspace)
        elif args.command == "render":
            print(f"rendered={render(args.workspace, args.pdf)}")
        elif args.command == "preprocess":
            print(f"preprocessed={preprocess(args.workspace)}")
        elif args.command == "merge":
            print(merge(args.workspace))
        elif args.command == "export":
            print(export_docx(args.workspace))
        else:
            cleanup(args.workspace, args.confirm)
    except (PipelineError, OSError, fitz.FileDataError) as error:
        print(f"error: {error}", file=sys.stderr)
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
